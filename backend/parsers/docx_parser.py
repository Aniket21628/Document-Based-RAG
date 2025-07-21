from docx import Document
from typing import Dict
import logging

logger = logging.getLogger(__name__)

class DOCXParser:
    @staticmethod
    def parse(file_path: str) -> Dict[str, any]:
        """Parse DOCX and extract text content"""
        try:
            doc = Document(file_path)
            
            text_content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text_content += row_text + "\n"
            
            metadata = {
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables),
                "file_type": "docx",
                "file_path": file_path
            }
            
            return {
                "content": text_content,
                "metadata": metadata
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise Exception(f"Failed to parse DOCX: {str(e)}")